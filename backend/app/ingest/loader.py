# backend/app/ingest/loader.py

from pathlib import Path
from typing import List, Iterator
from langchain_core.documents import Document as LangchainDocument
import hashlib
import logging

import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentLoader:
    """工业级文档加载器（无 Poppler / 高性能 / 适配 RAG）"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".html", ".md"}

    def __init__(self):
        pass

    def _generate_document_id(self, file_path: str) -> str:
        content = Path(file_path).read_bytes()
        return hashlib.sha256(content).hexdigest()[:16]

    # =========================
    # 📄 PDF 解析（核心优化）
    # =========================
    def _load_pdf(self, path: Path) -> List[LangchainDocument]:
        docs = []
        doc = fitz.open(str(path))

        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if not text:
                continue

            docs.append(LangchainDocument(
                page_content=text,
                metadata={
                    "source": path.name,
                    "filename": path.name,
                    "filetype": ".pdf",
                    "page_number": page_num,
                }
            ))

        return docs

    # =========================
    # 📄 DOCX 解析
    # =========================
    def _load_docx(self, path: Path) -> List[LangchainDocument]:
        doc = DocxDocument(str(path))
        docs = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            docs.append(LangchainDocument(
                page_content=text,
                metadata={
                    "source": path.name,
                    "filename": path.name,
                    "filetype": ".docx",
                    "paragraph": i
                }
            ))

        return docs

    # =========================
    # 📄 文本类文件
    # =========================
    def _load_text(self, path: Path) -> List[LangchainDocument]:
        text = path.read_text(encoding="utf-8", errors="ignore").strip()

        if not text:
            return []

        return [LangchainDocument(
            page_content=text,
            metadata={
                "source": path.name,
                "filename": path.name,
                "filetype": path.suffix.lower(),
            }
        )]

    # =========================
    # 🚀 主入口
    # =========================
    def load_file(self, file_path: str) -> List[LangchainDocument]:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()

        logger.info(f"📄 加载文件: {path.name}")

        try:
            if suffix == ".pdf":
                docs = self._load_pdf(path)

            elif suffix == ".docx":
                docs = self._load_docx(path)

            elif suffix in {".txt", ".md", ".html"}:
                docs = self._load_text(path)

            else:
                raise ValueError(f"不支持的文件格式: {suffix}")

            logger.info(f"✅ 加载完成: {len(docs)} chunks | {path.name}")
            return docs

        except Exception as e:
            logger.error(f"❌ 加载失败 {path.name}: {e}", exc_info=True)
            raise

    # =========================
    # 📁 批量加载
    # =========================
    def load_directory(self, directory: str, recursive: bool = True) -> Iterator[LangchainDocument]:
        path = Path(directory)

        for file_path in path.rglob("*") if recursive else path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    for doc in self.load_file(str(file_path)):
                        yield doc
                except Exception as e:
                    logger.warning(f"跳过文件 {file_path}: {e}")